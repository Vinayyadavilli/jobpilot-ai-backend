import io
import docx
from docx.shared import Pt


def generate_docx_from_text(enhanced_text: str) -> bytes:
    """
    Convert AI-enhanced markdown/text into a professional Word document (.docx).
    Formats headings, bullet points, and paragraphs cleanly.
    """
    doc = docx.Document()
    
    # Configure default professional font styling (Calibri/Arial style)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    lines = enhanced_text.splitlines()
    
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            # Add a bit of space between paragraph blocks
            continue
            
        # Parse basic markdown styles
        if stripped_line.startswith("# "):
            p = doc.add_heading(stripped_line[2:], level=1)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif stripped_line.startswith("## "):
            p = doc.add_heading(stripped_line[3:], level=2)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif stripped_line.startswith("### "):
            p = doc.add_heading(stripped_line[4:], level=3)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
        elif stripped_line.startswith("* ") or stripped_line.startswith("- "):
            # Bullet list items
            # Clean up potential leading asterisk or dash
            bullet_text = stripped_line[2:]
            # Remove bold marker if present (e.g. **Skill:** Text -> Skill: Text)
            bullet_text = bullet_text.replace("**", "")
            
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(3)
        else:
            # Normal paragraph text
            # Clean markdown bold markup
            paragraph_text = stripped_line.replace("**", "")
            
            p = doc.add_paragraph(paragraph_text)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            
    # Save document directly into an in-memory byte stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
