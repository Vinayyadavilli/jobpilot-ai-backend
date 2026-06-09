import io
import os
from pypdf import PdfReader
import docx


def extract_text_from_pdf(content: bytes) -> str:
    """Extract plain text from PDF file content bytes."""
    try:
        reader = PdfReader(io.BytesIO(content))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse PDF resume: {str(e)}")


def extract_text_from_docx(content: bytes) -> str:
    """Extract plain text from DOCX file content bytes (including tables)."""
    try:
        doc = docx.Document(io.BytesIO(content))
        text_parts = []
        
        # 1. Parse paragraphs
        for para in doc.paragraphs:
            if para.text:
                text_parts.append(para.text)
                
        # 2. Parse tables (common in resume formatting)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
                    
        return "\n".join(text_parts).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX resume: {str(e)}")


def extract_text(content: bytes, filename: str) -> str:
    """Determine file type by extension and extract its plain text."""
    ext = os.path.splitext(filename.lower())[1]
    
    if ext == ".pdf":
        return extract_text_from_pdf(content)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(content)
    elif ext in [".txt", ".md", ".rtf"]:
        return content.decode("utf-8", errors="ignore").strip()
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload a PDF, DOCX, TXT, MD, or RTF file.")
